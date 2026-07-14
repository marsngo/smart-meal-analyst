import streamlit as st
import requests
import pandas as pd
import io
import plotly.express as px
from google import genai
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- إعدادات الصفحة ---
st.set_page_config(page_title="محلل البيانات الإنسانية الذكي", page_icon="📊", layout="wide")

# --- واجهة المستخدم الذكية والمرنة ---
st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght=400;500;700&display=swap');
        html, body, [data-testid="stAppViewContainer"] {
            font-family: 'Tajawal', sans-serif;
            direction: RTL;
        }
        h1, h2, h3, h4, .stMarkdown p, label {
            text-align: right !important;
            direction: RTL !important;
            font-family: 'Tajawal', sans-serif !important;
        }
        [data-testid="id-textbox"], .stTextInput input, [data-testid="stFileUploader"] {
            direction: ltr !important;
            text-align: left !important;
        }
        .kpi-container {
            background-color: rgba(59, 130, 246, 0.08);
            border: 2px solid #3b82f6;
            padding: 20px;
            border-radius: 12px;
            text-align: center;
            margin-bottom: 15px;
        }
        .audit-card {
            background-color: rgba(239, 68, 68, 0.08);
            border-right: 5px solid #ef4444;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 15px;
        }
        .audit-card-warning {
            background-color: rgba(245, 158, 11, 0.08);
            border-right: 5px solid #f59e0b;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 15px;
        }
        .ai-narrative-box {
            background-color: rgba(16, 185, 129, 0.08);
            border-right: 5px solid #10b981;
            padding: 25px;
            border-radius: 10px;
            margin-top: 20px;
            margin-bottom: 25px;
        }
        .matrix-divider {
            border-top: 4px double #3b82f6;
            margin-top: 50px;
            margin-bottom: 30px;
        }
        .cross-section-divider {
            border-top: 2px dashed #4b5563;
            margin-top: 30px;
            margin-bottom: 30px;
        }
        .export-box {
            background-color: rgba(99, 102, 241, 0.08);
            border: 2px solid #6366f1;
            padding: 25px;
            border-radius: 12px;
            margin-top: 40px;
            text-align: center;
        }
        .rights-footer {
            margin-top: 80px;
            padding: 25px;
            background-color: rgba(49, 46, 129, 0.05);
            border-radius: 12px;
            text-align: center;
            border: 2px dashed #6366f1;
        }
        .rights-footer p, .rights-footer h4, .rights-footer span {
            text-align: center !important;
            direction: ltr !important;
        }
    </style>
""", unsafe_allow_html=True)

st.title("📊 منصة أدوات المسح الميداني والتحليل الذكي (Smart MEAL)")
st.subheader("إصدار مفتوح المصدر - نظام معالجة البيانات الإنسانية وصياغة التقارير الفورية")

if 'kobo_data' not in st.session_state:
    st.session_state.kobo_data = None
if 'generated_reports' not in st.session_state:
    st.session_state.generated_reports = []

embedded_api_key = ""
if "GEMINI_API_KEY" in st.secrets:
    embedded_api_key = st.secrets["GEMINI_API_KEY"]

with st.sidebar:
    st.markdown("### 🔑 إدارة مفتاح الذكاء الاصطناعي")
    if embedded_api_key:
        st.success("🔒 تم تفعيل الـ API السحابي المدمج بنجاح!")
        api_key = embedded_api_key
    else:
        api_key = st.text_input("أدخل مفتاح Gemini API يدوياً (اختياري):", type="password")
    
    st.markdown("---")
    st.markdown("### 🎨 تخصيص الهوية البصرية")
    color_theme = st.selectbox(
        "اختر ثيم الألوان للرسوم البيانية:",
        ("الأزرق القياسي (Streamlit)", "الألوان الإنسانية الدافئة (مجموعة قطاع الحماية)", "ثيم الطبيعة الحيوي (الأخضر والأزرق)", "ثيم اليونيسف الرسمي (UNICEF Theme)", "ألوان فخمة داكنة (Sunset)")
    )
    
    palette_map = {
        "الأزرق القياسي (Streamlit)": px.colors.qualitative.G10,
        "الألوان الإنسانية الدافئة (مجموعة قطاع الحماية)": px.colors.qualitative.Bold,
        "ثيم الطبيعة الحيوي (الأخضر والأزرق)": px.colors.qualitative.T10,
        "ثيم اليونيسف الرسمي (UNICEF Theme)": px.colors.qualitative.Set2,
        "ألوان فخمة داكنة (Sunset)": px.colors.qualitative.Dark24
    }
    selected_palette = palette_map[color_theme]

    st.markdown("---")
    st.markdown("### 🔌 إعدادات جلب البيانات")
    source_option = st.radio("اختر طريقة جلب البيانات الميدانية:", ("رفع ملف يدوي (Excel)", "سحب مباشر عبر الـ API"))
    
    if source_option == "سحب مباشر عبر الـ API":
        server_url = st.selectbox("اختر سيرفر الكوبو:", ("https://eu.kobotoolbox.org", "https://kf.kobotoolbox.org", "https://kobo.humanitarianresponse.info"))
        api_token = st.text_input("أدخل رمز الأمان (API Token):", type="password")
        asset_id = st.text_input("أدخل معرف الاستمارة (Asset ID):")
        if st.button("🚀 سحب البيانات الآن"):
            try:
                meta_url = f"{server_url}/api/v2/assets/{asset_id}/"
                headers = {"Authorization": f"Token {api_token}"}
                meta_response = requests.get(meta_url, headers=headers)
                
                label_mapping = {}
                if meta_response.status_code == 200:
                    meta_json = meta_response.json()
                    for item in meta_json.get('content', {}).get('survey', []):
                        if 'name' in item and 'label' in item:
                            lbl = item['label']
                            if isinstance(lbl, list) and len(lbl) > 0: label_mapping[item['name']] = str(lbl[0])
                            elif isinstance(lbl, dict): label_mapping[item['name']] = str(lbl.get('Arabic (ar)', list(lbl.values())[0]))
                            else: label_mapping[item['name']] = str(lbl)

                data_url = f"{server_url}/api/v2/assets/{asset_id}/data/?format=json"
                response = requests.get(data_url, headers=headers)
                if response.status_code == 200:
                    json_data = response.json()
                    if 'results' in json_data:
                        raw_df = pd.DataFrame(json_data['results'])
                        if label_mapping: raw_df.rename(columns=label_mapping, inplace=True)
                        cleaned_cols = [str(col).split('>')[-1].split('<')[0] if '<' in str(col) else str(col) for col in raw_df.columns]
                        final_cols = []
                        seen_cols = {}
                        for col in cleaned_cols:
                            if col.strip() == "" or col.lower() == "nan": col = "عمود_غير_معرف"
                            if col in seen_cols:
                                seen_cols[col] += 1
                                final_cols.append(f"{col}_{seen_cols[col]}")
                            else:
                                seen_cols[col] = 0
                                final_cols.append(col)
                        raw_df.columns = final_cols
                        st.session_state.kobo_data = raw_df
                        st.session_state.generated_reports = []
                        st.success("✅ تم سحب البيانات واستبدال الأكواد بالأسئلة العربية صراحة!")
                        st.rerun()
                else:
                    st.error(f"فشل الاتصال بكوبو. رمز الخطأ: {response.status_code}")
            except Exception as e:
                st.error(f"حدث خطأ أثناء معالجة العناوين والبيانات: {e}")

if source_option == "رفع ملف يدوي (Excel)":
    uploaded_file = st.file_uploader("قم برفع ملف البيانات الخام المستخرج من الكوبو (Excel):", type=["xlsx", "xls"])
    if uploaded_file is not None:
        try:
            raw_df = pd.read_excel(uploaded_file)
            cleaned_cols = [str(col).split('>')[-1].split('<')[0] if '<' in str(col) else str(col) for col in raw_df.columns]
            final_cols = []
            seen_cols = {}
            for col in cleaned_cols:
                if col.strip() == "" or col.lower() == "nan": col = "عمود_غير_معرف"
                if col in seen_cols:
                    seen_cols[col] += 1
                    final_cols.append(f"{col}_{seen_cols[col]}")
                else:
                    seen_cols[col] = 0
                    final_cols.append(col)
            raw_df.columns = final_cols
            st.session_state.kobo_data = raw_df
            st.success("✅ تم حفظ ملف الإكسل بنجاح!")
        except Exception as e:
            st.error(f"خطأ في قراءة الملف: {e}")

    if st.button("🗑️ مسح الجلسة وإعادة ضبط البيانات"):
        st.session_state.kobo_data = None
        st.session_state.generated_reports = []
        st.rerun()

# ==================== معالجة وعرض البيانات ====================
kobo_data = st.session_state.kobo_data

if kobo_data is not None:
    st.markdown("---")
    st.header("✨ لوحة القيادة التنفيذية وملخص المؤشرات (Executive Factsheet)")
    total_beneficiaries = kobo_data.shape[0]
    
    gender_col = [c for c in kobo_data.columns if 'جنس' in str(c).lower() or 'gender' in str(c).lower()]
    female_pct = "0%"
    if gender_col:
        female_count = kobo_data[kobo_data[gender_col[0]].astype(str).str.contains('أنثى|female|woman', case=False, na=False)].shape[0]
        female_pct = f"{(female_count / total_beneficiaries * 100):.1f}%" if total_beneficiaries > 0 else "0%"
        
    disability_col = [c for c in kobo_data.columns if 'disability' in str(c).lower() or 'إعاقة' in str(c).lower()]
    pwd_pct = "0%"
    if disability_col:
        pwd_count = kobo_data[kobo_data[disability_col[0]].astype(str).str.contains('نعم|yes|true|disabled', case=False, na=False)].shape[0]
        pwd_pct = f"{(pwd_count / total_beneficiaries * 100):.1f}%" if total_beneficiaries > 0 else "0%"

    kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
    with kpi_col1:
        st.markdown(f"""<div class="kpi-container"><h3>👥 إجمالي المستفيدين الكلي</h3><h1 style="font-size:42px; font-weight:700; margin-top:10px;">{total_beneficiaries:,}</h1></div>""", unsafe_allow_html=True)
    with kpi_col2:
        st.markdown(f"""<div class="kpi-container"><h3>👩‍🦰 نسبة الإناث في العينة</h3><h1 style="font-size:42px; font-weight:700; margin-top:10px;">{female_pct}</h1></div>""", unsafe_allow_html=True)
    with kpi_col3:
        st.markdown(f"""<div class="kpi-container"><h3>♿ نسبة إعاقات المستفيدين</h3><h1 style="font-size:42px; font-weight:700; margin-top:10px;">{pwd_pct}</h1></div>""", unsafe_allow_html=True)

    st.markdown("### 📋 نظرة عامة على البيانات المستلمة")
    st.dataframe(kobo_data.head())
    
    # --- قسم مقارنة القبلي والبعدي المطور الذكي ---
    st.markdown("---")
    st.header("🔄 قسم مقارنة التقييم القبلي والبعدي (Pre vs Post Assessment)")
    
    analysis_type = st.radio("اختر هيكلية تصميم ملف البيانات الحالي لديك:", ("الملف يحتوي على عمودين منفصلين (عمود للقبلي وعمود للبعدي)", "الملف يحتوي على عمود سكور واحد + عمود نصي يحدد نوع التقييم (قبلي/بعدي)"))
    
    mean_pre, mean_post = 0.0, 0.0
    calculation_ready = False
    paired_df = None
    
    if analysis_type == "الملف يحتوي على عمودين منفصلين (عمود للقبلي وعمود للبعدي)":
        pre_post_col1, pre_post_col2 = st.columns(2)
        with pre_post_col1:
            pre_column = st.selectbox("🎯 اختر عمود التقييم القبلي (Pre-test score):", ["لا يوجد"] + list(kobo_data.columns))
        with pre_post_col2:
            post_column = st.selectbox("🎯 اختر عمود التقييم البعدي (Post-test score):", ["لا يوجد"] + list(kobo_data.columns))
            
        if pre_column != "لا يوجد" and post_column != "لا يوجد":
            pre_numeric = pd.to_numeric(kobo_data[pre_column], errors='coerce')
            post_numeric = pd.to_numeric(kobo_data[post_column], errors='coerce')
            mean_pre = pre_numeric.mean()
            mean_post = post_numeric.mean()
            calculation_ready = True
            
    else:
        filter_col1, filter_col2, filter_col3 = st.columns(3)
        with filter_col1:
            user_id_column = st.selectbox("🆔 اختر عمود اسم المستفيد أو الـ ID الفريد:", ["لا يوجد"] + list(kobo_data.columns))
        with filter_col2:
            score_column = st.selectbox("📊 اختر العمود الذي يحتوي على الدرجة والأسكور (Total Score):", ["لا يوجد"] + list(kobo_data.columns))
        with filter_col3:
            type_column = st.selectbox("🏷️ اختر العمود الذي يحدد نوع التقييم (مثال: نوع التقييم):", ["لا يوجد"] + list(kobo_data.columns))
            
        if score_column != "لا يوجد" and type_column != "لا يوجد" and user_id_column != "لا يوجد":
            pre_rows = kobo_data[kobo_data[type_column].astype(str).str.contains('قبلي|pre', case=False, na=False)].copy()
            post_rows = kobo_data[kobo_data[type_column].astype(str).str.contains('بعدي|post', case=False, na=False)].copy()
            
            if not pre_rows.empty and not post_rows.empty:
                pre_rows[score_column] = pd.to_numeric(pre_rows[score_column], errors='coerce')
                post_rows[score_column] = pd.to_numeric(post_rows[score_column], errors='coerce')
                
                paired_df = pd.merge(
                    pre_rows[[user_id_column, score_column]], 
                    post_rows[[user_id_column, score_column]], 
                    on=user_id_column, 
                    suffixes=('_قبلي_Pre', '_بعدي_Post')
                )
                
                if not paired_df.empty:
                    mean_pre = paired_df[f'{score_column}_قبلي_Pre'].mean()
                    mean_post = paired_df[f'{score_column}_بعدي_Post'].mean()
                    paired_df['الفارق الصافي والتحسن (Impact)'] = paired_df[f'{score_column}_بعدي_Post'] - paired_df[f'{score_column}_قبلي_Pre']
                    calculation_ready = True
                else:
                    st.warning("⚠️ فشل الربط التلقائي. تأكد من أن أسماء المستفيدين مكتوبة بنفس الأحرف والاملاء تماماً في المقابلتين القبلية والبعدية.")
            else:
                st.warning("⚠️ لم نجد الكلمات الدلالية ('قبلي' أو 'بعدي') داخل العمود النصي المختار لتصفية السطور.")

    if calculation_ready and not pd.isna(mean_pre) and not pd.isna(mean_post):
        improvement = mean_post - mean_pre
        
        comp_col1, comp_col2, comp_col3 = st.columns(3)
        with comp_col1:
            st.metric(label="📊 متوسط القبلي العام (Pre Mean)", value=f"{mean_pre:.2f}")
        with comp_col2:
            st.metric(label="📈 متوسط البعدي العام (Post Mean)", value=f"{mean_post:.2f}")
        with comp_col3:
            st.metric(label="✨ الفارق الصافي والتحسن الإجمالي", value=f"{improvement:+.2f}", delta=f"{improvement:.2f}")
            
        df_compare = pd.DataFrame({
            'المرحلة (Assessment)': ['التقييم القبلي (Pre)', 'التقييم البعدي (Post)'],
            'متوسط الدرجة الكلية (Mean Score)': [mean_pre, mean_post]
        })
        fig_compare = px.bar(df_compare, x='المرحلة (Assessment)', y='متوسط الدرجة الكلية (Mean Score)', text='متوسط الدرجة الكلية (Mean Score)', title="مقارنة الأداء العام بين القبلي والبعدي", color='المرحلة (Assessment)', color_discrete_sequence=selected_palette)
        fig_compare.update_traces(texttemplate='%{text:.2f}', textposition='outside')
        st.plotly_chart(fig_compare, use_container_width=True)
        
        if paired_df is not None:
            st.markdown("### 👤 جدول قياس الفارق الفردي لكل مستفيد (Individual Impact Tracking)")
            st.dataframe(paired_df, use_container_width=True)
            
    elif calculation_ready:
        st.error("⚠️ فشل حساب الأرقام. يرجى التحقق من أن عمود الدرجة يحتوي على أرقام صافية وخالٍ من النصوص.")

    # --- مصفوفة التحليل الكبرى للفئات والأعمدة الديموغرافية ---
    st.markdown("---")
    st.header("🎛️ مصفوفة التحليل والمقارنات العامة (Bulk Analysis)")
    analysis_col1, analysis_col2 = st.columns(2)
    with analysis_col1:
        main_questions = st.multiselect("1. اختر الأسئلة والمؤشرات المراد تحليلها نسب تكرارية:", list(kobo_data.columns))
    with analysis_col2:
        cross_questions = st.multiselect("2. اختر محاور التقاطع والديموغرافيا (مثال: الجنس، المحافظة):", list(kobo_data.columns))

    st.markdown("### 🎯 إعدادات تخصيص صياغة التقرير")
    donor_selection = st.selectbox(
        "اختر المانح المستهدف لتوجه التقرير له:",
        ("صياغة عامة رصينة", "OCHA (تركيز صارم على الأرقام، الاختصار، والاحتياج العاجل)", "UNICEF (تركيز على حماية الأطفال، الدمج، والنوع الاجتماعي)", "ECHO / SIDA (تركيز على الكفاءة، معايير المساءلة، والاستدامة)")
    )

    if main_questions:
        run_ai_global = False
        if cross_questions:
            st.markdown("### 🤖 التحكم بتقرير الذكاء الاصطناعي الشامل")
            if not api_key:
                st.info("💡 لتوليد تقارير سردية تلقائية عبر Gemini، يرجى تفعيل الـ API Key أولاً.")
            else:
                run_ai_global = st.toggle("⚙️ تشغيل خبير الـ MEAL الذكي لإنتاج التقارير السردية تلقائياً")

            if st.button("🔄 بدء معالجة وتحديث مصفوفة الألوان والمانح الحالي"):
                st.session_state.generated_reports = []
                st.rerun()

        for current_main in main_questions:
            st.markdown(f'<div class="matrix-divider"></div>', unsafe_allow_html=True)
            st.markdown(f"## 📊 كتلة التحليل الأساسية للسؤال: [{current_main}]")
            
            st.markdown("🔍 **خيارات العرض الفوري للرسم البياني الحالي:**")
            chart_selector_col, style_col = st.columns([1, 2])
            with chart_selector_col:
                chart_type = st.selectbox(f"اختر نوع الرسم لـ [{current_main[:15]}...]:", ("أعمدة رأسي (Bar)", "دائرة (Pie)", "دونات (Donut)", "خطوط (Line)"), key=f"type_{current_main}")

            if not cross_questions:
                df_counts = kobo_data[current_main].value_counts().reset_index()
                df_counts.columns = [current_main, 'العدد']
                df_counts['النسبة المئوية (%)'] = (df_counts['العدد'] / df_counts['العدد'].sum() * 100).round(1)
                
                table_side_col, chart_side_col = st.columns([1, 2])
                with table_side_col:
                    st.write("📋 جدول النسب التكرارية:")
                    st.dataframe(df_counts, use_container_width=True)
                with chart_side_col:
                    if chart_type == "أعمدة رأسي (Bar)":
                        fig = px.bar(df_counts, x=current_main, y='النسبة المئوية (%)', text=df_counts['النسبة المئوية (%)'].astype(str) + '%', title=f"التوزيع المئوي لإجابات: {current_main}", color=current_main, color_discrete_sequence=selected_palette)
                        fig.update_traces(textposition='outside')
                    elif chart_type == "دائرة (Pie)":
                        fig = px.pie(df_counts, values='العدد', names=current_main, color=current_main, title=f"التوزيع المئوي لإجابات: {current_main}", color_discrete_sequence=selected_palette)
                    elif chart_type == "دونات (Donut)":
                        fig = px.pie(df_counts, values='العدد', names=current_main, color=current_main, hole=0.4, title=f"التوزيع المئوي لإجابات: {current_main}", color_discrete_sequence=selected_palette)
                    else:  # Line Chart
                        fig = px.line(df_counts, x=current_main, y='النسبة المئوية (%)', markers=True, title=f"التوزيع المئوي لإجابات: {current_main}", color_discrete_sequence=selected_palette)
                    
                    st.plotly_chart(fig, use_container_width=True)
            else:
                for current_cross in cross_questions:
                    st.markdown(f'<div class="cross-section-divider"></div>', unsafe_allow_html=True)
                    st.markdown(f"### 🔗 المقارنة المئوية المتقاطعة: [{current_main}] ✖️ [{current_cross}]")
                    
                    crosstab_pct = pd.crosstab(kobo_data[current_main], kobo_data[current_cross], normalize='columns') * 100
                    crosstab_pct = crosstab_pct.round(1)
                    df_melted_pct = pd.melt(crosstab_pct.reset_index(), id_vars=[current_main], value_name='النسبة المئوية (%)')
                    
                    table_side_col, chart_side_col = st.columns([1, 2])
                    with table_side_col:
                        st.write("📋 جدول النسب المئوية المتقاطعة (%):")
                        st.dataframe(crosstab_pct.map(lambda x: f"{x}%"), use_container_width=True)
                    with chart_side_col:
                        if chart_type == "أعمدة رأسي (Bar)":
                            fig = px.bar(df_melted_pct, x=current_main, y='النسبة المئوية (%)', color=current_cross, barmode='group', text=df_melted_pct['النسبة المئوية (%)'].astype(str) + '%', title=f"المقارنة المئوية لـ [{current_main}] حسب فئات [{current_cross}]", color_discrete_sequence=selected_palette)
                            fig.update_traces(textposition='outside')
                        elif chart_type == "دائرة (Pie)":
                            fig = px.pie(df_melted_pct, values='النسبة المئوية (%)', names=current_main, color=current_main, title=f"المقارنة المئوية لـ [{current_main}] حسب فئات [{current_cross}]", color_discrete_sequence=selected_palette)
                        elif chart_type == "دونات (Donut)":
                            fig = px.pie(df_melted_pct, values='النسبة المئوية (%)', names=current_main, color=current_main, hole=0.4, title=f"المقارنة المئوية لـ [{current_main}] حسب فئات [{current_cross}]", color_discrete_sequence=selected_palette)
                        else:  # Line Chart
                            fig = px.line(df_melted_pct, x=current_main, y='النسبة المئوية (%)', color=current_cross, markers=True, title=f"المقارنة المئوية لـ [{current_main}] حسب فئات [{current_cross}]", color_discrete_sequence=selected_palette)
                        
                        st.plotly_chart(fig, use_container_width=True)

                    narrative_text = ""
                    if run_ai_global and api_key:
                        existing = [r for r in st.session_state.generated_reports if r['main'] == current_main and r['cross'] == current_cross and r['donor'] == donor_selection and r['theme'] == color_theme]
                        if not existing:
                            with st.spinner(f"جاري صياغة التقرير السردي لمحور [{current_cross}]..."):
                                try:
                                    client = genai.Client(api_key=api_key)
                                    donor_instructions = "اكتب بأسلوب تقارير المنظمات الدولية الفخم والعام."
                                    if "OCHA" in donor_selection:
                                        donor_instructions = "صِغ النص بأسلوب أوتشا (OCHA): ركز بشدة على الاختصار البليغ وإبراز الأرقام وتوضيح الفجوة في الاحتياج العاجل."
                                    elif "UNICEF" in donor_selection:
                                        donor_instructions = "صِغ النص بأسلوب اليونيسف (UNICEF): ركز بقوة على مؤشرات حماية الأطفال، الفئات الأكثر ضعفاً، وأبعاد النوع الاجتماعي."
                                    elif "ECHO" in donor_selection:
                                        donor_instructions = "صِغ النص بأسلوب إيكو (ECHO / SIDA): ركز على معايير الكفاءة والمساءلة الإنسانية، واستدامة الخدمة."

                                    prompt = f"""أنت مستشار أول لتقارير المتابعة والتقييم (MEAL Specialist) في منظمات دولية.
                                    اكتب فقرة تحليلية سردية احترافية باللغة العربية الفصحى للتقرير النهائي بناءً على جدول النسب المئوية التالي:
                                    السؤال الرئيسي: {current_main}
                                    محور التفكيك: {current_cross}
                                    جدول النسب المئوية:
                                    {crosstab_pct.to_string()}
                                    ⚠️ شرط الصياغة الخاص بالمانح: {donor_instructions}
                                    قم بتحليل الأرقام بدقة بالغة وبدون مقدمات تقليدية، واختم بتوصية ميدانية واحدة ذكية مبنية على الجدول."""
                                    
                                    response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                                    narrative_text = response.text

                                    st.session_state.generated_reports.append({
                                        'main': current_main,
                                        'cross': current_cross,
                                        'donor': donor_selection,
                                        'theme': color_theme,
                                        'text': narrative_text,
                                        'image': None
                                    })
                                except Exception as e:
                                    st.error(f"خطأ في الـ AI: {e}")
                                    narrative_text = "فشل توليد التقرير"
                        else:
                            narrative_text = existing[0]['text']

                        st.markdown(f"""<div class="ai-narrative-box"><h4>📝 التقرير السردي المخصص لـ [{donor_selection}]:</h4><div style="line-height:1.8; font-size:15px; text-align:justify;">{narrative_text.replace('\n', '<br>')}</div></div>""", unsafe_allow_html=True)

            # --- مركز التحميل الكلي لملفات الوورد ---
            if 'generated_reports' in st.session_state and st.session_state.generated_reports:
                valid_reports = [r for r in st.session_state.generated_reports if r['donor'] == donor_selection and r['theme'] == color_theme]
                if valid_reports:
                    st.markdown('<div class="export-box">', unsafe_allow_html=True)
                    st.markdown("### 📥 مركز تحميل المخرجات والتقارير الموجهة الملونة")
                    doc = Document()
                    for section in doc.sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)

                    doc.add_heading(f"التقرير التحليلي المخصص للمانح ({donor_selection}) - ثيم {color_theme}", level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    for report in valid_reports:
                        h2 = doc.add_heading(f"📊 المتغير: {report['main']}", level=2)
                        h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        h3 = doc.add_heading(f"🔗 محور التفكيك: {report['cross']}", level=3)
                        h3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        p = doc.add_paragraph(report['text'])
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.line_spacing = 1.5
                        doc.add_page_break()

                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    st.download_button(label=f"📥 تحميل تقرير الـ Word الملون الكامل (.docx)", data=bio, file_name=f"Comprehensive_Report_{donor_selection.split(' ')[0]}_{color_theme.split(' ')[0]}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    st.markdown('</div>', unsafe_allow_html=True)
else:
    st.info("💡 بانتظار ربط البيانات للبدء بالتحليل التنفيذي والمتقدم...")

# ==================== 💎 تذييل حفظ الحقوق الملكية ====================
st.markdown("""
    <div class="rights-footer">
        <h4 style="margin:0; font-weight:700; font-size:18px; color: #312e81 !important;">💡 Smart MEAL Platform | منصة التحليل الذكي للمتابعة والتقييم</h4>
        <p style="margin:8px 0 0 0; font-size:14px; color: #4b5563 !important;">Designed & Developed by: <b>Mhedy Alkhaldi</b> © 2026</p>
        <p style="margin:5px 0 0 0; font-size:12px; color: #6b7280 !important; font-style:italic;">Open-Source Humanitarian Solution - Dedicated for General Benefit & Philanthropy</p>
    </div>
""", unsafe_allow_html=True)
